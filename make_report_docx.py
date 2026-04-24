"""Convert report.md to report.docx preserving all content and structure."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Colour palette ───────────────────────────────────────────────────────────
INDIGO  = RGBColor(0x63, 0x66, 0xf1)   # headings
DARK    = RGBColor(0x1e, 0x20, 0x3a)   # title
MUTED   = RGBColor(0x44, 0x4d, 0x5e)   # body text
CODE_BG = RGBColor(0xf3, 0xf4, 0xf6)   # code block shading
CODE_FG = RGBColor(0x1f, 0x29, 0x37)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def shade_paragraph(para, hex_color="F3F4F6"):
    """Add background shading to a paragraph (for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, "Calibri", 22, bold=True, color=DARK)
    elif level == 2:
        set_font(run, "Calibri", 16, bold=True, color=INDIGO)
    else:
        set_font(run, "Calibri", 13, bold=True, color=INDIGO)
    return p

def add_body(doc, text):
    """Add a body paragraph, handling **bold** and `code` inline marks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

    # Split on **bold** and `code` markers
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=MUTED)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9, color=CODE_FG)
            run.font.highlight_color = None
        else:
            run = p.add_run(part)
            set_font(run, color=MUTED)
    return p

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        shade_paragraph(p, "F0F0F0")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run(line)
        set_font(run, name="Courier New", size=9, color=CODE_FG)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)

    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=MUTED)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9, color=CODE_FG)
        else:
            run = p.add_run(part)
            set_font(run, color=MUTED)
    return p

def add_numbered(doc, text, n):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run_num = p.add_run(f"{n}. ")
    set_font(run_num, bold=True, color=INDIGO)

    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=MUTED)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9, color=CODE_FG)
        else:
            run = p.add_run(part)
            set_font(run, color=MUTED)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(h)
        set_font(run, bold=True, size=9, color=RGBColor(0xff, 0xff, 0xff))
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '6366F1')
        tcPr.append(shd)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        fill = 'F8F9FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cells[ci].text = ''
            pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
            parts = re.split(pattern, str(cell_text))
            p = cells[ci].paragraphs[0]
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, bold=True, size=9, color=MUTED)
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Courier New", size=8, color=CODE_FG)
                else:
                    run = p.add_run(part)
                    set_font(run, size=9, color=MUTED)
            tc = cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    doc.add_paragraph()  # spacer after table

# ════════════════════════════════════════════════════════════════════════════
# Parse and render report.md
# ════════════════════════════════════════════════════════════════════════════

with open("report.md", "r") as f:
    lines = f.readlines()

in_code = False
code_lines = []
in_table = False
table_headers = []
table_rows = []
numbered_counter = 0
i = 0

while i < len(lines):
    line = lines[i].rstrip('\n')

    # ── Code block ──────────────────────────────────────────────────────────
    if line.startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            if code_lines:
                add_code_block(doc, code_lines)
                doc.add_paragraph()
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # ── Table ────────────────────────────────────────────────────────────────
    if line.startswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        if not in_table:
            in_table = True
            table_headers = cells
            table_rows = []
            i += 1
            continue
        if re.match(r'^[\s\-\|]+$', line):
            i += 1
            continue
        table_rows.append(cells)
        i += 1
        # peek: if next line is not a table row, flush
        if i >= len(lines) or not lines[i].startswith('|'):
            add_table(doc, table_headers, table_rows)
            in_table = False
        continue

    # ── Headings ─────────────────────────────────────────────────────────────
    if line.startswith('### '):
        numbered_counter = 0
        add_heading(doc, line[4:], 3)
    elif line.startswith('## '):
        numbered_counter = 0
        add_heading(doc, line[3:], 2)
    elif line.startswith('# '):
        numbered_counter = 0
        add_heading(doc, line[2:], 1)

    # ── Horizontal rule ──────────────────────────────────────────────────────
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C7D2FE')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Bullets ──────────────────────────────────────────────────────────────
    elif line.startswith('- '):
        numbered_counter = 0
        add_bullet(doc, line[2:])

    # ── Numbered list ─────────────────────────────────────────────────────────
    elif re.match(r'^\d+\. ', line):
        m = re.match(r'^(\d+)\. (.*)', line)
        if m:
            add_numbered(doc, m.group(2), int(m.group(1)))

    # ── Empty line ────────────────────────────────────────────────────────────
    elif line.strip() == '':
        pass  # skip blank lines (spacing handled per-element)

    # ── Body text ─────────────────────────────────────────────────────────────
    else:
        text = line.strip()
        if text:
            add_body(doc, text)

    i += 1

doc.save("report.docx")
print("Saved: report.docx")
